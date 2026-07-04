"""
pdf_export.py — PDF, Word ve CSV dışa aktarma fonksiyonları
"""

import io
import json
import csv
from datetime import datetime
from typing import Optional


# --------------------------------------------------------------------------- #
#  PDF Dışa Aktarma                                                             #
# --------------------------------------------------------------------------- #

def generate_pdf(sinav: dict, meta: dict) -> bytes:
    """
    Sınav içeriğini PDF olarak oluşturur ve bytes döndürür.
    
    Args:
        sinav: AI'ın ürettiği sınav içeriği dict'i
        meta: {konu, zorluk, soru_sayisi, tarih} bilgileri
    """
    try:
        from fpdf import FPDF
    except ImportError:
        raise ImportError("fpdf2 yüklü değil: pip install fpdf2")

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # Unicode desteği için sistem fontunu yüklemeyi dene
    font_family = "Helvetica"
    has_custom_font = False
    
    try:
        import os
        # Windows sistem fontları yolu
        win_font_dir = r"C:\Windows\Fonts"
        normal_font = os.path.join(win_font_dir, "arial.ttf")
        bold_font = os.path.join(win_font_dir, "arialbd.ttf")
        
        if os.path.exists(normal_font) and os.path.exists(bold_font):
            pdf.add_font("ArialUnicode", "", normal_font)
            pdf.add_font("ArialUnicode", "B", bold_font)
            font_family = "ArialUnicode"
            has_custom_font = True
    except Exception:
        pass

    def safe_text(text: str) -> str:
        """Eğer özel font yüklenememişse PDF uyumlu ASCII benzeri metin döndürür."""
        if has_custom_font:
            return text
        replacements = {
            'ş': 's', 'Ş': 'S', 'ı': 'i', 'İ': 'I',
            'ğ': 'g', 'Ğ': 'G', 'ü': 'u', 'Ü': 'U',
            'ö': 'o', 'Ö': 'O', 'ç': 'c', 'Ç': 'C',
        }
        for tr, en in replacements.items():
            text = text.replace(tr, en)
        return text

    # Başlık
    pdf.set_font(font_family, "B", 20)
    pdf.set_fill_color(67, 97, 238)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(0, 14, safe_text("AI Sınav Hazırlayıcı"), fill=True, ln=True, align="C")
    pdf.ln(4)

    # Meta bilgileri
    pdf.set_font(font_family, "", 11)
    pdf.set_text_color(50, 50, 50)
    pdf.set_fill_color(240, 242, 255)
    meta_text = (
        f"Konu: {meta.get('konu', '')}   |   "
        f"Zorluk: {meta.get('zorluk', '')}   |   "
        f"Soru Sayısı: {meta.get('soru_sayisi', '')}   |   "
        f"Tarih: {meta.get('tarih', '')}"
    )
    pdf.multi_cell(0, 8, safe_text(meta_text), fill=True, align="C")
    pdf.ln(6)

    def section_title(title: str):
        pdf.set_font(font_family, "B", 13)
        pdf.set_text_color(67, 97, 238)
        pdf.cell(0, 10, safe_text(title), ln=True)
        pdf.set_draw_color(67, 97, 238)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(3)
        pdf.set_text_color(30, 30, 30)
        pdf.set_font(font_family, "", 11)

    # 1. Konu Özeti
    section_title("1. Konu Özeti")
    pdf.multi_cell(0, 7, safe_text(sinav.get("ozet", "")))
    pdf.ln(5)

    # 2. Test Soruları
    section_title("2. Test Soruları")
    for i, soru_obj in enumerate(sinav.get("test_sorulari", []), 1):
        pdf.set_font(font_family, "B", 11)
        pdf.multi_cell(0, 7, safe_text(f"Soru {i}: {soru_obj.get('soru', '')}"))
        pdf.set_font(font_family, "", 11)
        secenekler = soru_obj.get("secenekler", {})
        if isinstance(secenekler, dict):
            for harf, metin in secenekler.items():
                pdf.set_x(pdf.l_margin + 8)
                w = pdf.w - pdf.x - pdf.r_margin
                pdf.multi_cell(w, 6, safe_text(f"{harf}) {metin}"))
        pdf.ln(3)

    # 3. Cevap Anahtarı
    section_title("3. Cevap Anahtarı")
    answers = []
    for i, soru_obj in enumerate(sinav.get("test_sorulari", []), 1):
        answers.append(f"Soru {i}: {soru_obj.get('dogru', '?')}")
    pdf.multi_cell(0, 7, safe_text("   |   ".join(answers)))
    pdf.ln(5)

    # 4. Klasik Sorular
    section_title("4. Klasik Sorular")
    for i, soru in enumerate(sinav.get("klasik_sorular", []), 1):
        pdf.multi_cell(0, 7, safe_text(f"{i}. {soru}"))
        pdf.ln(2)

    # 5. Çalışma Önerisi
    section_title("5. Çalışma Önerisi")
    pdf.multi_cell(0, 7, safe_text(sinav.get("calisma_onerisi", "")))

    return bytes(pdf.output())


# --------------------------------------------------------------------------- #
#  Word Dışa Aktarma                                                            #
# --------------------------------------------------------------------------- #

def generate_word(sinav: dict, meta: dict) -> bytes:
    """
    Sınav içeriğini Word (.docx) olarak oluşturur ve bytes döndürür.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx yüklü değil: pip install python-docx")

    doc = Document()

    # Başlık
    title = doc.add_heading("AI Sınav Hazırlayıcı", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Meta bilgiler
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta_para.add_run(
        f"Konu: {meta.get('konu', '')}  |  "
        f"Zorluk: {meta.get('zorluk', '')}  |  "
        f"Soru Sayısı: {meta.get('soru_sayisi', '')}  |  "
        f"Tarih: {meta.get('tarih', '')}"
    )
    run.bold = True

    doc.add_paragraph()

    def add_section(title_text: str, level: int = 1):
        h = doc.add_heading(title_text, level=level)
        return h

    # 1. Konu Özeti
    add_section("1. Konu Özeti")
    doc.add_paragraph(sinav.get("ozet", ""))

    # 2. Test Soruları
    add_section("2. Test Soruları")
    for i, soru_obj in enumerate(sinav.get("test_sorulari", []), 1):
        p = doc.add_paragraph()
        run = p.add_run(f"Soru {i}: {soru_obj.get('soru', '')}")
        run.bold = True

        secenekler = soru_obj.get("secenekler", {})
        if isinstance(secenekler, dict):
            for harf, metin in secenekler.items():
                doc.add_paragraph(f"    {harf}) {metin}", style="List Bullet")
        doc.add_paragraph()

    # 3. Cevap Anahtarı
    add_section("3. Cevap Anahtarı")
    for i, soru_obj in enumerate(sinav.get("test_sorulari", []), 1):
        doc.add_paragraph(
            f"Soru {i}: {soru_obj.get('dogru', '?')}",
            style="List Number",
        )

    # 4. Klasik Sorular
    add_section("4. Klasik Sorular")
    for i, soru in enumerate(sinav.get("klasik_sorular", []), 1):
        doc.add_paragraph(f"{i}. {soru}")

    # 5. Çalışma Önerisi
    add_section("5. Çalışma Önerisi")
    doc.add_paragraph(sinav.get("calisma_onerisi", ""))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# --------------------------------------------------------------------------- #
#  CSV Dışa Aktarma                                                             #
# --------------------------------------------------------------------------- #

def generate_csv(sinav: dict, meta: dict) -> str:
    """
    Sınav sorularını CSV formatında döndürür.
    """
    output = io.StringIO()
    writer = csv.writer(output)

    # Başlık satırı
    writer.writerow([
        "Soru No", "Soru", "A Seçeneği", "B Seçeneği",
        "C Seçeneği", "D Seçeneği", "Doğru Cevap"
    ])

    for i, soru_obj in enumerate(sinav.get("test_sorulari", []), 1):
        secenekler = soru_obj.get("secenekler", {})
        writer.writerow([
            i,
            soru_obj.get("soru", ""),
            secenekler.get("A", "") if isinstance(secenekler, dict) else "",
            secenekler.get("B", "") if isinstance(secenekler, dict) else "",
            secenekler.get("C", "") if isinstance(secenekler, dict) else "",
            secenekler.get("D", "") if isinstance(secenekler, dict) else "",
            soru_obj.get("dogru", ""),
        ])

    return output.getvalue()
